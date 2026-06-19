"""
layer1_ingestion.py  —  Layer 1: File Ingestion
─────────────────────────────────────────────────
Reads resume files and returns raw text strings.

Supported formats:
  - PDF   →  pdfminer.six
  - DOCX  →  python-docx  (paragraphs + tables)
  - TXT   →  plain file read
  - MD    →  plain file read (markdown treated as plain text)

Changes in this version:
  - BUG FIX (_extract_docx): Now reads BOTH doc.paragraphs AND doc.tables.
    Previously only paragraphs were read, so any resume built with DOCX
    tables (very common — skills grids, experience tables, academic detail
    tables) was returning almost no text to Layer 2.
  - Merged cell deduplication: DOCX merged cells repeat the same text
    across multiple cell objects. We track seen cell content and skip
    duplicates so skills don't appear 3-4x in the extracted text.

Usage:
  from layer1_ingestion import load_resume
  text = load_resume("resumes/john_doe.pdf")
"""

import os
from pathlib import Path

# ── PDF ───────────────────────────────────────────────────────────────────────
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfparser import PDFSyntaxError

# ── DOCX ──────────────────────────────────────────────────────────────────────
from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def load_resume(file_path: str) -> str:
    """
    Extract raw text from a resume file.

    Args:
        file_path: Absolute or relative path to the resume file.

    Returns:
        Raw text as a single string.

    Raises:
        ValueError: If the file format is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()

    extractors = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".txt":  _extract_txt,
        ".md":   _extract_txt,
    }

    if ext not in extractors:
        raise ValueError(
            f"Unsupported format '{ext}'. "
            f"Supported: {list(extractors.keys())}"
        )

    raw_text = extractors[ext](path)

    if len(raw_text.strip()) < 50:
        print(f"[Warning] Very little text extracted from '{path.name}'. "
              f"File may be scanned/image-based (OCR not supported yet).")

    return raw_text


def load_resumes_from_folder(folder_path: str) -> dict[str, str]:
    """
    Load all supported resume files from a folder.

    Args:
        folder_path: Path to folder containing resume files.

    Returns:
        Dict mapping filename → extracted text.
    """
    folder = Path(folder_path)
    if not folder.is_dir():
        raise NotADirectoryError(f"Not a directory: {folder_path}")

    supported_extensions = {".pdf", ".docx", ".txt", ".md"}
    results = {}

    for file in sorted(folder.iterdir()):
        if file.suffix.lower() not in supported_extensions:
            continue
        try:
            results[file.name] = load_resume(str(file))
            print(f"[✓] Loaded: {file.name}")
        except Exception as e:
            print(f"[✗] Skipped '{file.name}': {e}")

    print(f"\nLoaded {len(results)} resume(s) from '{folder_path}'")
    return results


# ─────────────────────────────────────────────────────────────────────────────
# Private extractors
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfminer.six."""
    try:
        text = pdf_extract_text(str(path))
        return text or ""
    except PDFSyntaxError as e:
        raise ValueError(f"Could not parse PDF '{path.name}': {e}") from e


def _extract_docx(path: Path) -> str:
    """
    Extract text from a DOCX file — reads both paragraphs AND tables.

    BUG FIX: The original implementation only read doc.paragraphs.
    Many resumes (especially Indian engineering resume templates) use
    DOCX tables for skills, academic details, projects, and internship
    sections. python-docx stores table content in doc.tables, not
    doc.paragraphs, so those sections returned zero text before this fix.

    Merged cell handling: DOCX merged cells repeat the same cell object
    multiple times across rows/columns. We deduplicate by tracking seen
    cell text to avoid repeating skills/content 3-4x.
    """
    doc = Document(str(path))
    parts = []

    # ── 1. Regular paragraphs (headers, summary, contact info) ───────────────
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    # ── 2. Table cells (skills, experience, education, projects, etc.) ────────
    # Track seen content to handle merged cells (they repeat the same text
    # across multiple cell objects — deduplicate to avoid inflating skill counts)
    seen_cells: set[str] = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt and txt not in seen_cells:
                    seen_cells.add(txt)
                    parts.append(txt)

    return "\n".join(parts)


def _extract_txt(path: Path) -> str:
    """Read a plain text or markdown file."""
    for encoding in ("utf-8", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode '{path.name}' with utf-8 or latin-1.")


# ─────────────────────────────────────────────────────────────────────────────
# Quick local test  (python layer1_ingestion.py)
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python layer1_ingestion.py <path_to_resume>")
        sys.exit(1)

    text = load_resume(sys.argv[1])
    print("─" * 60)
    print(f"Extracted {len(text)} characters")
    print("─" * 60)
    print(text[:1000], "..." if len(text) > 1000 else "")